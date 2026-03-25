import pytest
from pathlib import Path

from claw.processors.document import extract_text_file, extract_docx, DocResult, MAX_FILE_SIZE


class TestExtractTextFile:
    def test_reads_txt(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("Hello world this is test content", encoding="utf-8")
        result = extract_text_file(f)
        assert result.success is True
        assert result.word_count == 6
        assert "Hello" in result.text

    def test_file_not_found(self, tmp_path):
        result = extract_text_file(tmp_path / "nope.txt")
        assert result.success is False

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        result = extract_text_file(f)
        assert result.success is True
        assert result.word_count == 0

    def test_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin.txt"
        f.write_bytes("caf\xe9 world".encode("latin-1"))
        result = extract_text_file(f)
        assert result.success is True
        assert "caf" in result.text

    def test_title_from_filename(self, tmp_path):
        f = tmp_path / "my-notes.txt"
        f.write_text("content", encoding="utf-8")
        result = extract_text_file(f)
        assert result.title == "my-notes"

    def test_oversized_file(self, tmp_path):
        f = tmp_path / "big.txt"
        f.write_bytes(b"x" * (MAX_FILE_SIZE + 1))
        result = extract_text_file(f)
        assert result.success is False

    def test_result_immutable(self):
        r = DocResult(text="t", word_count=1, title="T", file_type="txt", success=True, error="")
        with pytest.raises(AttributeError):
            r.text = "x"


class TestExtractDocx:
    def test_file_not_found(self, tmp_path):
        result = extract_docx(tmp_path / "nope.docx")
        assert result.success is False

    def test_valid_docx(self, tmp_path):
        from docx import Document
        f = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello world this is a test")
        doc.add_paragraph("Second paragraph here")
        doc.save(str(f))

        result = extract_docx(f)
        assert result.success is True
        assert result.word_count > 0
        assert "Hello" in result.text
        assert result.file_type == "docx"

    def test_oversized_docx(self, tmp_path):
        f = tmp_path / "big.docx"
        f.write_bytes(b"x" * (MAX_FILE_SIZE + 1))
        result = extract_docx(f)
        assert result.success is False
